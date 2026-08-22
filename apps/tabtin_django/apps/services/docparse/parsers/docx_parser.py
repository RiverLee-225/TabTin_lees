"""
Word 文档解析器 (v0.4)

v0.4 改进：
- DC-11: 段落/内联样式保真——按 run 提取 bold/italic/underline/strike/
  superscript/subscript/文字颜色/高亮/字符样式（Strong/Emphasis 等），
  以及段落对齐方式，供下游直接构建 pmJson marks（不再经纯 markdown 中转丢样式）

v0.3 改进：
- DC-7: 图片 blob 通过 base64 持久化到 metadata，支持下游上传 OSS
- DC-8: 不再依赖 doc.paragraphs 索引对齐，使用元素→对象映射
- DC-10: 识别 Code 样式和等宽字体段落，标记为 codeBlock

v0.2 改进：
- 分页符检测：识别 <w:br w:type="page"/> 进行逻辑分页
- 中文标题样式支持：匹配 "标题 1"、"标题 2" 等中文/日文/韩文样式名
- 图片提取：从 paragraph 的 inline shapes 中提取嵌入图片
- 更丰富的列表检测
"""

from __future__ import annotations

import base64
import hashlib
import logging
import os
import re
import zipfile
from io import BytesIO
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn

from .base import BaseDocumentParser, ChunkResult, PageResult, ParseResult
from .legacy_doc_convert import (
    UNRECOGNIZED_WORD_PAYLOAD_MSG,
    convert_word_payload_to_docx,
    detect_word_payload_kind,
    extract_rtf_plaintext,
)
from .registry import register_parser

logger = logging.getLogger(__name__)

MAX_DOCX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB on-disk
MAX_DOCX_UNCOMPRESSED_SIZE = 200 * 1024 * 1024  # 200 MB total decompressed
MAX_DOCX_COMPRESSION_RATIO = 50  # flag zip-bombs with ratio > 50:1
MAX_IMAGE_BLOB_SIZE = 5 * 1024 * 1024  # 5 MB per image for base64 persistence


@register_parser
class DocxParser(BaseDocumentParser):

    def supported_mimes(self) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]

    def parse(self, file_path: str, **kwargs) -> ParseResult:
        converted_path: str | None = None
        try:
            kind = detect_word_payload_kind(file_path)
            if kind == "docx_zip":
                _validate_docx_safe(file_path)
                return self._parse_docx(file_path)

            if kind in ("ole", "rtf", "html"):
                try:
                    logger.info("检测到 Word 兼容载荷 kind=%s，尝试转换为 .docx", kind)
                    converted_path = convert_word_payload_to_docx(file_path)
                    _validate_docx_safe(converted_path)
                    return self._parse_docx(converted_path)
                except ValueError:
                    if kind == "rtf":
                        logger.warning("RTF→docx 转换失败，降级为纯文本提取")
                        return _parse_rtf_as_plaintext(file_path)
                    if kind == "html":
                        logger.warning("HTML→docx 转换失败，降级为 HTML 文本解析")
                        from .plaintext_parser import PlaintextParser

                        return PlaintextParser().parse(file_path)
                    raise

            raise ValueError(UNRECOGNIZED_WORD_PAYLOAD_MSG)
        finally:
            if converted_path:
                try:
                    os.unlink(converted_path)
                except OSError:
                    pass

    def _parse_docx(self, file_path: str) -> ParseResult:
        doc = DocxDocument(file_path)

        page_chunks: list[list[ChunkResult]] = [[]]
        current_page = 0
        seq = 0
        title = ""

        # DC-8: build element→object maps to avoid index misalignment
        # when textboxes / headers / SDT wrappers cause doc.paragraphs
        # to diverge from body element order.
        _para_by_elem = {id(p._element): p for p in doc.paragraphs}
        _tbl_by_elem = {id(t._element): t for t in doc.tables}

        # Refs : numId/ilvl → bullet|ordered + 层级，供 _classify_paragraph
        # 判定真正的列表项（而不是弱启发式）。
        numbering_defs = _load_numbering_definitions(doc)

        for elem in doc.element.body:
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

            if tag == "p":
                para = _para_by_elem.get(id(elem))
                if para is None:
                    continue

                if _has_page_break(para):
                    current_page += 1
                    page_chunks.append([])
                    seq = 0

                images = _extract_images_from_paragraph(para, doc)
                for img_info in images:
                    seq += 1
                    img_meta: dict = {
                        "source": "structural",
                        "image_hash": img_info["hash"],
                        "content_type": img_info["content_type"],
                    }
                    if img_info.get("image_b64"):
                        img_meta["image_b64"] = img_info["image_b64"]
                    page_chunks[current_page].append(ChunkResult(
                        chunk_type="image",
                        content=f"[嵌入图片: {img_info['filename']}]",
                        sequence=seq,
                        bbox=None,
                        metadata=img_meta,
                    ))

                text = para.text.strip()
                if not text:
                    continue

                list_info = _resolve_list_info(numbering_defs, para)
                chunk_type, heading_level, list_meta = _classify_paragraph(para, list_info)
                seq += 1

                if not title and chunk_type == "heading" and heading_level and heading_level <= 2:
                    title = text[:200]

                is_bold = any(run.bold for run in para.runs if run.bold is not None)
                font_size = None
                for run in para.runs:
                    if run.font and run.font.size:
                        font_size = run.font.size.pt
                        break

                metadata: dict = {
                    "source": "structural",
                    "bold": is_bold,
                }
                if font_size:
                    metadata["font_size"] = font_size
                if list_meta:
                    metadata.update(list_meta)

                links = _extract_hyperlinks(para)
                if links:
                    metadata["links"] = links

                # DC-11: heading/paragraph/list 走 run 级样式保真——由
                # service._docx_rich_pm_node / 列表节点构建直接消费 pmJson marks，
                # 绕开 markdown 中转（颜色/高亮/上下标等无法用纯 markdown 表达）。
                # codeBlock 暂不接入。
                if chunk_type in ("heading", "paragraph", "list"):
                    runs_meta = _extract_run_marks(para)
                    # Refs : 启发式列表（无 numPr，靠样式名/缩进识别）文本里
                    # 带着字面量符号（"• "/"1. "），真正转成 listItem 后交由编辑器
                    # 渲染标记，这里去掉前缀避免符号重复；numPr 列表本身不含该前缀。
                    if chunk_type == "list" and list_meta and list_meta.get("list_num_id") is None:
                        runs_meta = _strip_list_marker_from_runs(runs_meta, text)
                    if runs_meta:
                        metadata["runs"] = runs_meta
                    align = _paragraph_align(para)
                    if align:
                        metadata["align"] = align

                page_chunks[current_page].append(ChunkResult(
                    chunk_type=chunk_type,
                    content=text,
                    sequence=seq,
                    bbox=None,
                    heading_level=heading_level,
                    metadata=metadata,
                ))

            elif tag == "tbl":
                table = _tbl_by_elem.get(id(elem))
                if table is None:
                    continue

                md = _table_to_markdown(table)
                if not md:
                    continue

                seq += 1
                page_chunks[current_page].append(ChunkResult(
                    chunk_type="table",
                    content=md,
                    sequence=seq,
                    bbox=None,
                    metadata={
                        "source": "structural",
                        "rows": len(table.rows),
                        "cols": len(table.columns),
                    },
                ))

        # 构建 pages
        pages: list[PageResult] = []
        for page_num, chunks in enumerate(page_chunks, 1):
            if not chunks:
                continue
            text_content = "\n".join(c.content for c in chunks if c.content)
            pages.append(PageResult(
                page_number=page_num,
                width=0,
                height=0,
                chunks=chunks,
                text_content=text_content,
            ))

        if not pages:
            pages = [PageResult(page_number=1, width=0, height=0, chunks=[], text_content="")]

        return ParseResult(
            pages=pages,
            title=title,
            parse_method="structural",
        )


# ======================================================================
# 标题样式匹配（中文 + 英文 + 多语言）
# ======================================================================

_HEADING_PATTERNS = [
    re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE),
    re.compile(r"^标题\s*(\d+)$"),
    re.compile(r"^見出し\s*(\d+)$"),       # 日文
    re.compile(r"^제목\s*(\d+)$"),         # 韩文
    re.compile(r"^Titre\s*(\d+)$", re.IGNORECASE),  # 法文
    re.compile(r"^Überschrift\s*(\d+)$", re.IGNORECASE),  # 德文
]

_TITLE_PREFIXES = ("Title", "标题", "見出し")
_SUBTITLE_PREFIXES = ("Subtitle", "副标题")
_LIST_PREFIXES = ("List", "Bullet", "列表", "项目符号")

# DC-10: code block style / font detection
_CODE_STYLE_PATTERNS = [
    re.compile(r"^Code$", re.IGNORECASE),
    re.compile(r"^Code\s*(Block|Char)$", re.IGNORECASE),
    re.compile(r"^(HTML|Source)\s*Code$", re.IGNORECASE),
    re.compile(r"^Verbatim", re.IGNORECASE),
    re.compile(r"^PlainText$", re.IGNORECASE),
    re.compile(r"^程序代码$"),
    re.compile(r"^代码"),
]

_MONOSPACE_FONTS = frozenset({
    "courier new", "consolas", "monaco", "menlo",
    "source code pro", "fira code", "roboto mono",
    "ubuntu mono", "dejavu sans mono", "liberation mono",
    "lucida console", "andale mono",
})


# ======================================================================
# Refs : 编号列表识别——读取 w:numPr (numId/ilvl) + numbering.xml，
# 区分 bullet / ordered 并保留层级，供 service._build_tabdoc_import_draft
# 构建真正的 nested bulletList/orderedList（而不是把列表项当纯段落）。
# ======================================================================


def _load_numbering_definitions(doc) -> dict[str, dict[int, dict[str, Any]]]:
    """解析 numbering.xml：numId(str) -> {ilvl(int): {"kind", "start"}}。

    numFmt == "bullet" 判定为 bullet，"none" 表示该级别不显示编号（视为非列表），
    其余数字/字母/罗马数字等格式统一归为 ordered。lvlOverride/startOverride
    （某个 num 实例对某一级别的起始编号覆盖）优先于 abstractNum 的默认 start。
    """
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        return {}
    if numbering_part is None:
        return {}
    numbering_elm = numbering_part.element

    abstract_levels: dict[str, dict[int, dict[str, Any]]] = {}
    for abstract_num in numbering_elm.findall(qn("w:abstractNum")):
        abstract_id = abstract_num.get(qn("w:abstractNumId"))
        if abstract_id is None:
            continue
        levels: dict[int, dict[str, Any]] = {}
        for lvl in abstract_num.findall(qn("w:lvl")):
            ilvl_raw = lvl.get(qn("w:ilvl"))
            if ilvl_raw is None:
                continue
            fmt_elem = lvl.find(qn("w:numFmt"))
            num_fmt = fmt_elem.get(qn("w:val")) if fmt_elem is not None else None
            start_elem = lvl.find(qn("w:start"))
            try:
                start = int(start_elem.get(qn("w:val"))) if start_elem is not None else 1
            except (TypeError, ValueError):
                start = 1
            if num_fmt == "bullet":
                kind = "bullet"
            elif num_fmt == "none":
                kind = "none"
            else:
                kind = "ordered"
            levels[int(ilvl_raw)] = {"kind": kind, "start": start}
        abstract_levels[abstract_id] = levels

    result: dict[str, dict[int, dict[str, Any]]] = {}
    for num_elem in numbering_elm.findall(qn("w:num")):
        num_id = num_elem.get(qn("w:numId"))
        abstract_ref = num_elem.find(qn("w:abstractNumId"))
        if num_id is None or abstract_ref is None:
            continue
        levels = dict(abstract_levels.get(abstract_ref.get(qn("w:val")), {}))
        for override in num_elem.findall(qn("w:lvlOverride")):
            ilvl_raw = override.get(qn("w:ilvl"))
            start_elem = override.find(qn("w:startOverride"))
            if ilvl_raw is None or start_elem is None:
                continue
            try:
                ilvl = int(ilvl_raw)
                start = int(start_elem.get(qn("w:val")))
            except (TypeError, ValueError):
                continue
            level_info = dict(levels.get(ilvl, {"kind": "ordered"}))
            level_info["start"] = start
            levels[ilvl] = level_info
        result[num_id] = levels
    return result


def _paragraph_num_pr(para) -> tuple[str, int] | None:
    """段落生效的 (numId, ilvl)：先看段落自身 pPr/numPr，再回退到段落样式的
    pPr/numPr（沿 basedOn 链向上找）——Word 内置 "List Bullet"/"List Number"
    等样式把编号定义挂在样式上而非段落本身，是最常见的落地方式。
    numId == "0" 表示显式移除编号（不是列表）。
    """
    num_pr = None
    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        num_pr = pPr.find(qn("w:numPr"))
    if num_pr is None:
        style = para.style
        seen_styles: set[int] = set()
        while style is not None and id(style.element) not in seen_styles:
            seen_styles.add(id(style.element))
            style_pPr = style.element.find(qn("w:pPr"))
            if style_pPr is not None:
                num_pr = style_pPr.find(qn("w:numPr"))
                if num_pr is not None:
                    break
            style = getattr(style, "base_style", None)
    if num_pr is None:
        return None

    num_id_elem = num_pr.find(qn("w:numId"))
    if num_id_elem is None:
        return None
    num_id = num_id_elem.get(qn("w:val"))
    if not num_id or num_id == "0":
        return None

    ilvl_elem = num_pr.find(qn("w:ilvl"))
    try:
        ilvl = int(ilvl_elem.get(qn("w:val"))) if ilvl_elem is not None else 0
    except (TypeError, ValueError):
        ilvl = 0
    return num_id, max(0, ilvl)


def _resolve_list_info(
    numbering_defs: dict[str, dict[int, dict[str, Any]]], para
) -> dict[str, Any] | None:
    """结合 numbering_defs 判定段落是否为编号列表项。

    返回 ``{"kind": "bullet"|"ordered", "level": int, "num_id": str, "start": int|None}``，
    非列表项（无 numPr，或该级别 numFmt=="none"）返回 None。
    """
    num_pr = _paragraph_num_pr(para)
    if num_pr is None:
        return None
    num_id, ilvl = num_pr
    level_info = numbering_defs.get(num_id, {}).get(ilvl)
    if level_info is None or level_info.get("kind") == "none":
        return None
    return {
        "kind": level_info["kind"],
        "level": ilvl,
        "num_id": num_id,
        "start": level_info.get("start"),
    }


def _classify_paragraph(
    para, list_info: dict[str, Any] | None = None
) -> tuple[str, int | None, dict[str, Any] | None]:
    """返回 (chunk_type, heading_level, list_meta)。

    list_meta 非空时 chunk_type == "list"，携带
    ``{"list_kind", "list_level", "list_num_id", "list_start"}``，供
    service._build_tabdoc_import_draft 构建真正的 nested bulletList/orderedList
    （Refs ）。list_info 来自 numbering.xml（真实 numId/ilvl），优先于
    样式名/缩进的弱启发式判定。
    """
    style_name = para.style.name if para.style else ""

    # DC-10: detect code blocks before headings (a "Code" style should not
    # accidentally match heading patterns)
    if _is_code_style(style_name) or _is_monospace_paragraph(para):
        return "codeBlock", None, None

    for pattern in _HEADING_PATTERNS:
        m = pattern.match(style_name)
        if m:
            return "heading", int(m.group(1)), None

    if any(style_name.startswith(p) for p in _TITLE_PREFIXES):
        return "heading", 1, None

    if any(style_name.startswith(p) for p in _SUBTITLE_PREFIXES):
        return "heading", 2, None

    if list_info is not None:
        return "list", None, {
            "list_kind": list_info["kind"],
            "list_level": list_info["level"],
            "list_num_id": list_info["num_id"],
            "list_start": list_info["start"],
        }

    if any(style_name.startswith(p) for p in _LIST_PREFIXES):
        return "list", None, {
            "list_kind": _heuristic_list_kind(para.text),
            "list_level": 0,
            "list_num_id": None,
            "list_start": None,
        }

    if para.paragraph_format and para.paragraph_format.left_indent:
        indent_pt = para.paragraph_format.left_indent.pt if para.paragraph_format.left_indent else 0
        if indent_pt > 20:
            text = para.text.strip()
            if text and (text[0] in "•-–—·※◆" or re.match(r"^\d+[.)]\s", text)):
                return "list", None, {
                    "list_kind": _heuristic_list_kind(text),
                    "list_level": 0,
                    "list_num_id": None,
                    "list_start": None,
                }

    return "paragraph", None, None


def _heuristic_list_kind(text: str) -> str:
    """启发式（无 numPr）列表种类判定：数字/字母前缀视为 ordered，否则 bullet。"""
    stripped = text.strip()
    if stripped and re.match(r"^\d+[.)]\s", stripped):
        return "ordered"
    return "bullet"


def _is_code_style(style_name: str) -> bool:
    return any(p.match(style_name) for p in _CODE_STYLE_PATTERNS)


def _is_monospace_paragraph(para) -> bool:
    """All runs with an explicit font must use a monospace font family."""
    if para.style and hasattr(para.style, "font") and para.style.font:
        pstyle_font = para.style.font.name
        if pstyle_font and pstyle_font.lower() in _MONOSPACE_FONTS:
            return True
    if not para.runs:
        return False
    mono_count = 0
    total_with_font = 0
    for run in para.runs:
        font_name = run.font.name if run.font else None
        if font_name:
            total_with_font += 1
            if font_name.lower() in _MONOSPACE_FONTS:
                mono_count += 1
    return total_with_font > 0 and mono_count == total_with_font


# ======================================================================
# 分页符检测
# ======================================================================

def _has_page_break(para) -> bool:
    """检测段落是否包含分页符 <w:br w:type="page"/>"""
    for run in para.runs:
        for br in run._element.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True

    # 也检查段落前的分页符（段落属性中的 pageBreakBefore）
    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        pb = pPr.find(qn("w:pageBreakBefore"))
        if pb is not None and pb.get(qn("w:val"), "true") != "false":
            return True

    return False


# ======================================================================
# 图片提取
# ======================================================================

def _image_result_from_rel(doc, embed_id: str) -> dict | None:
    """Resolve a document relationship id to an image payload dict."""
    try:
        rel = doc.part.rels.get(embed_id)
        if rel is None:
            return None
        img_part = rel.target_part
        img_bytes = img_part.blob
        result = {
            "filename": img_part.partname.split("/")[-1],
            "content_type": img_part.content_type,
            "hash": hashlib.md5(img_bytes).hexdigest()[:12],
        }
        if len(img_bytes) <= MAX_IMAGE_BLOB_SIZE:
            result["image_b64"] = base64.b64encode(img_bytes).decode("ascii")
        else:
            logger.info(
                "图片过大 (%d bytes)，跳过 base64 持久化: %s",
                len(img_bytes), result["filename"],
            )
        return result
    except Exception as exc:
        logger.debug("图片提取失败: %s", exc)
        return None


def _extract_images_from_paragraph(para, doc) -> list[dict]:
    """从段落中提取嵌入图片的基本信息和 blob 数据。

    支持 DrawingML（``w:drawing`` / ``a:blip``）以及旧式 VML（``w:pict`` /
    ``v:imagedata``）。WPS/LibreOffice 把 .doc 转成 .docx 时，内嵌图常仍是 VML。

    DC-7: blob 以 base64 编码存入返回字典的 ``image_b64`` 键，
    下游可据此上传 OSS。超过 MAX_IMAGE_BLOB_SIZE 的图片仅保留 hash。
    """
    images: list[dict] = []
    seen: set[str] = set()

    def _append(embed_id: str | None) -> None:
        if not embed_id or embed_id in seen:
            return
        result = _image_result_from_rel(doc, embed_id)
        if result is None:
            return
        seen.add(embed_id)
        images.append(result)

    # Traverse the whole paragraph tree so VML picts / drawings are found even
    # when they are not exposed via python-docx's ``paragraph.runs``.
    for element in para._element.iter():
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "blip":
            _append(element.get(qn("r:embed")))
        elif tag == "imagedata":
            # VML uses r:id (not r:embed)
            _append(element.get(qn("r:id")))
    return images


# ======================================================================
# 超链接提取
# ======================================================================

def _extract_hyperlinks(para) -> list[dict]:
    """提取段落中的超链接"""
    links = []
    for hyperlink in para._element.findall(qn("w:hyperlink")):
        r_id = hyperlink.get(qn("r:id"))
        if not r_id:
            continue
        try:
            rel = para.part.rels.get(r_id)
            if rel and hasattr(rel, "target_ref"):
                text = "".join(
                    node.text or ""
                    for node in hyperlink.findall(".//" + qn("w:t"))
                )
                links.append({"text": text, "url": rel.target_ref})
        except Exception:
            pass
    return links


# ======================================================================
# DC-11: 段落对齐 + run 级内联样式提取（保真 bold/italic/underline/strike/
# 上下标/文字颜色/高亮/字符样式），供 service._docx_rich_pm_node 直接构建
# pmJson marks，不经过会丢样式的纯 markdown 中转层。
# ======================================================================

_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
}
_JUSTIFY_ALIGN_VALUES = frozenset({
    WD_ALIGN_PARAGRAPH.JUSTIFY,
    WD_ALIGN_PARAGRAPH.JUSTIFY_LOW,
    WD_ALIGN_PARAGRAPH.JUSTIFY_MED,
    WD_ALIGN_PARAGRAPH.JUSTIFY_HI,
    WD_ALIGN_PARAGRAPH.DISTRIBUTE,
})

# Word/WPS 内置字符样式（w:rStyle）——这些样式本身承载格式（斜体/加粗/主题色），
# 但不一定会同时写直接 run 属性，因此需要按样式名兜底映射，否则「subtle /
# strong / intense emphasis」这类样式会被判定为无格式的纯文本。
_CHAR_STYLE_MARKS: dict[str, dict[str, Any]] = {
    "strong": {"bold": True},
    "emphasis": {"italic": True},
    "subtle emphasis": {"italic": True, "color": "#8C8C8C"},
    "intense emphasis": {"bold": True, "italic": True, "color": "#4472C4"},
}

# WD_COLOR_INDEX 是离散调色板（非任意 RGB），按名称映射到可直接写入 CSS 的颜色值。
_HIGHLIGHT_COLOR_MAP: dict[str, str] = {
    "YELLOW": "#fef9c3",
    "BRIGHT_GREEN": "#00FF00",
    "GREEN": "#dcfce7",
    "TURQUOISE": "turquoise",
    "TEAL": "teal",
    "BLUE": "#dbeafe",
    "DARK_BLUE": "#00008B",
    "PINK": "#fce7f3",
    "RED": "#fee2e2",
    "DARK_RED": "#8B0000",
    "VIOLET": "#f3e8ff",
    "DARK_YELLOW": "#808000",
    "GRAY_25": "#f3f4f6",
    "GRAY_50": "#808080",
    "BLACK": "black",
    "WHITE": "white",
}

_HEX_COLOR_RE = re.compile(r"^[0-9A-Fa-f]{6}$")

# 主题色（w:themeColor）→ 近似 hex 的静态兜底表，按 Office 2013+ 默认主题
# （dk1/lt1/dk2/lt2/accent1-6/hyperlink）取色。仅在 run 没有随附显式 RGB
# （见 _run_color_hex）时启用——低成本弱近似，自定义过配色的主题文档会不准，
# 因为 python-docx 不提供解析文档实际 theme1.xml 的公开 API，这里没有读取
# 该文档真实的主题配色定义。
_THEME_COLOR_APPROX: dict[str, str] = {
    "dark1": "#000000",
    "text1": "#000000",
    "light1": "#FFFFFF",
    "background1": "#FFFFFF",
    "dark2": "#44546A",
    "text2": "#44546A",
    "light2": "#E7E6E6",
    "background2": "#E7E6E6",
    "accent1": "#4472C4",
    "accent2": "#ED7D31",
    "accent3": "#A5A5A5",
    "accent4": "#FFC000",
    "accent5": "#5B9BD5",
    "accent6": "#70AD47",
    "hyperlink": "#0563C1",
    "followedHyperlink": "#954F72",
}


def _style_chain(style):
    """沿 base_style 链逐级向上遍历样式（含自身），命中环引用即止。

    Word 的段落/字符样式支持链式继承（如「标题 1」可能 basedOn「标题基础」），
    对齐方式、下划线等属性若某一级样式没有直接设置，需要继续往上一级样式找。
    """
    seen: set[int] = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        yield style
        try:
            style = style.base_style
        except Exception:
            break


def _style_chain_alignment(style):
    """沿样式链找第一个显式设置了 paragraph_format.alignment 的样式。"""
    for s in _style_chain(style):
        try:
            pf = s.paragraph_format
        except Exception:
            continue
        if pf is None:
            continue
        try:
            alignment = pf.alignment
        except Exception:
            alignment = None
        if alignment is not None:
            return alignment
    return None


def _paragraph_align(para) -> str | None:
    """段落对齐方式 → pmJson textAlign 取值。

    para.alignment 是段落直接格式；很多标题「居中」其实写在样式定义
    （style.paragraph_format.alignment）而非每个标题段落的直接格式，
    直接格式为空时需要沿 base_style 链回退到样式设置，否则标题对齐会丢。
    """
    try:
        alignment = para.alignment
    except Exception:
        alignment = None
    if alignment is None:
        try:
            alignment = _style_chain_alignment(para.style)
        except Exception:
            alignment = None
    if alignment is None:
        return None
    if alignment in _ALIGN_MAP:
        return _ALIGN_MAP[alignment]
    if alignment in _JUSTIFY_ALIGN_VALUES:
        return "justify"
    return None


def _run_char_style_marks(run) -> dict[str, Any]:
    """按 run 引用的字符样式名（如 Strong/Emphasis）取默认格式兜底。"""
    try:
        style = run.style
        name = (style.name or "").strip().lower() if style else ""
    except Exception:
        name = ""
    base = _CHAR_STYLE_MARKS.get(name)
    return dict(base) if base else {}


def _color_hex_from_color_format(color) -> str | None:
    """ColorFormat → ``#RRGGBB``；显式 RGB 优先，其次主题色弱近似。"""
    if color is None:
        return None
    try:
        if color.type is None:
            return None
        rgb = color.rgb
        if rgb is not None:
            return f"#{rgb}"
        theme_color = getattr(color, "theme_color", None)
        if theme_color is not None:
            xml_value = getattr(theme_color, "xml_value", None)
            if xml_value:
                return _THEME_COLOR_APPROX.get(xml_value)
    except Exception:
        return None
    return None


def _style_chain_color_hex(style) -> str | None:
    """沿样式链找第一个显式字色（标题蓝字常写在 Heading 样式而非每个 run）。"""
    for s in _style_chain(style):
        font = getattr(s, "font", None)
        if font is None:
            continue
        try:
            hex_color = _color_hex_from_color_format(font.color)
        except Exception:
            hex_color = None
        if hex_color:
            return hex_color
    return None


def _run_color_hex(run, para_style=None) -> str | None:
    """run 的文字颜色：优先取显式 RGB；没有显式 RGB 但引用了主题色
    （w:themeColor）时，按 _THEME_COLOR_APPROX 做弱近似兜底；再回退
    字符样式链 / 段落样式链（calibre 标题蓝字写在 Heading 1/2 样式上）。

    真实 Word 保存主题色时通常会同时写一份当前快照 RGB（.rgb 可直接读到，
    见 python-docx ColorFormat.rgb 文档），因此这里的近似表主要覆盖
    「只写了 themeColor、没写快照 RGB」的边界场景（如部分第三方工具生成
    的 docx）；命中真实快照 RGB 时不会走近似表。
    """
    try:
        direct = _color_hex_from_color_format(run.font.color)
    except Exception:
        direct = None
    if direct:
        return direct

    try:
        char_style = run.style
    except Exception:
        char_style = None
    inherited = _style_chain_color_hex(char_style) if char_style is not None else None
    if inherited is None and para_style is not None:
        inherited = _style_chain_color_hex(para_style)
    return inherited


def _run_highlight_color(run) -> str | None:
    """run 的高亮色：优先取 Word 高亮调色板，其次回退到字符底纹（w:shd fill）。

    字符底纹回退主要用于覆盖「反色/自定义底色」一类不走标准高亮调色板的样式——
    这是启发式兜底，不是通用反色（inverse video）支持，见 DC-11 已知限制。
    """
    try:
        wd_color = run.font.highlight_color
    except Exception:
        wd_color = None
    if wd_color is not None and wd_color != WD_COLOR_INDEX.AUTO:
        mapped = _HIGHLIGHT_COLOR_MAP.get(wd_color.name)
        if mapped:
            return mapped
    try:
        rpr = run._element.find(qn("w:rPr"))
        shd = rpr.find(qn("w:shd")) if rpr is not None else None
        if shd is not None:
            fill = shd.get(qn("w:fill"))
            if fill and _HEX_COLOR_RE.match(fill) and fill.upper() != "FFFFFF":
                return f"#{fill}"
    except Exception:
        pass
    return None


def _iter_paragraph_runs(para) -> list[tuple[Any, str | None]]:
    """按文档顺序遍历段落内的 run，包括超链接内部的 run（python-docx 的
    ``paragraph.runs`` 不包含 ``<w:hyperlink>`` 内的 run，会漏掉链接文字）。

    返回 ``(Run, hyperlink_url_or_None)`` 序列。
    """
    from docx.text.run import Run as DocxRun

    results: list[tuple[Any, str | None]] = []
    for child in para._element:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "r":
            results.append((DocxRun(child, para), None))
        elif tag == "hyperlink":
            r_id = child.get(qn("r:id"))
            href = None
            if r_id:
                try:
                    rel = para.part.rels.get(r_id)
                    if rel is not None and hasattr(rel, "target_ref"):
                        href = rel.target_ref
                except Exception:
                    href = None
            for r_elem in child.findall(qn("w:r")):
                results.append((DocxRun(r_elem, para), href))
    return results


def _style_chain_underline(style):
    """沿样式链找第一个显式设置了 font.underline 的样式。"""
    for s in _style_chain(style):
        font = getattr(s, "font", None)
        if font is None:
            continue
        try:
            underline = font.underline
        except Exception:
            underline = None
        if underline is not None:
            return underline
    return None


def _run_underline(run, para_style=None) -> bool:
    """判断 run 最终是否带下划线。

    优先级：run 直接格式（run.underline）> run 引用的字符样式链
    （w:rStyle，超链接常见的 Hyperlink 字符样式多在这一级带下划线）
    > 段落样式链（w:pStyle 的默认 run 属性）。

    已知限制：不读取 `<w:docDefaults>` 里的全局默认 rPr——python-docx
    未提供该层的公开访问入口，因此「整份文档默认下划线」这一级不会被
    继承到；这类设置极少见（几乎没有文档在 docDefaults 里开全局下划线）。
    """
    direct = run.underline
    if direct is not None:
        return bool(direct)

    try:
        char_style = run.style
    except Exception:
        char_style = None
    inherited = _style_chain_underline(char_style) if char_style is not None else None

    if inherited is None and para_style is not None:
        inherited = _style_chain_underline(para_style)

    return bool(inherited) if inherited is not None else False


def _run_marks(run, para_style=None) -> dict[str, Any]:
    """单个 run 的格式集合：字符样式兜底 + 直接格式覆盖（直接格式优先级更高）。"""
    marks = _run_char_style_marks(run)

    if run.bold:
        marks["bold"] = True
    elif run.bold is False:
        marks.pop("bold", None)

    if run.italic:
        marks["italic"] = True
    elif run.italic is False:
        marks.pop("italic", None)

    if _run_underline(run, para_style):
        marks["underline"] = True

    font = run.font
    if font.strike or font.double_strike:
        marks["strike"] = True
    if font.superscript:
        marks["superscript"] = True
    elif font.subscript:
        marks["subscript"] = True

    color = _run_color_hex(run, para_style)
    if color:
        marks["color"] = color

    highlight = _run_highlight_color(run)
    if highlight:
        marks["highlight"] = highlight

    return marks


def _extract_run_marks(para) -> list[dict[str, Any]]:
    """把整段落按 run 拆成 ``[{"text": ..., <marks>}, ...]``，供下游直接
    构建 pmJson text node + marks，不需要经过丢样式的 markdown 中转。
    """
    try:
        para_style = para.style
    except Exception:
        para_style = None

    runs_meta: list[dict[str, Any]] = []
    for run, href in _iter_paragraph_runs(para):
        text = run.text
        if not text:
            continue
        entry: dict[str, Any] = {"text": text}
        entry.update(_run_marks(run, para_style))
        if href:
            entry["link"] = href
        runs_meta.append(entry)
    return runs_meta


# Refs : 启发式（无 numPr）列表段落文本里带着字面量符号（"• "/"1. "），
# 匹配前缀字符集需与 _classify_paragraph 的缩进+符号判定保持一致。
_LIST_MARKER_RE = re.compile(r"^[•\-–—·※◆]\s*|^\d+[.)]\s*")


def _strip_list_marker_from_runs(
    runs_meta: list[dict[str, Any]], full_text: str
) -> list[dict[str, Any]]:
    """去掉启发式列表段落开头的字面量符号，真正转成 listItem 后交由编辑器渲染
    标记，避免符号重复。仅从最前面的 run 裁剪，符号刚好跨多个 run 边界的极端
    情况不覆盖（已知限制）。
    """
    m = _LIST_MARKER_RE.match(full_text)
    if not m:
        return runs_meta
    remaining = len(m.group(0))
    adjusted: list[dict[str, Any]] = []
    for entry in runs_meta:
        entry_text = entry.get("text", "")
        if remaining <= 0:
            adjusted.append(entry)
            continue
        if len(entry_text) <= remaining:
            remaining -= len(entry_text)
            continue
        new_entry = dict(entry)
        new_entry["text"] = entry_text[remaining:]
        remaining = 0
        adjusted.append(new_entry)
    return adjusted


def _parse_rtf_as_plaintext(file_path: str) -> ParseResult:
    """RTF 无转换器时的纯文本降级，保证导入不因 DOCX/ZIP 校验直接失败。"""
    text = extract_rtf_plaintext(file_path)
    first_line = text.split("\n", 1)[0].strip() if text else ""
    title = first_line[:200] if first_line else ""
    content = text or "(空 RTF 文件)"
    return ParseResult(
        pages=[
            PageResult(
                page_number=1,
                width=0,
                height=0,
                chunks=[
                    ChunkResult(
                        chunk_type="paragraph",
                        content=content,
                        sequence=1,
                    )
                ],
                text_content=content,
            )
        ],
        title=title,
        parse_method="rtf_text_fallback",
    )


# ======================================================================
# 表格转 Markdown
# ======================================================================

def _table_to_markdown(table) -> str:
    rows_data: list[list[str]] = []
    for row in table.rows:
        cells = [
            cell.text.replace("|", "\\|").replace("\n", " ").strip()
            for cell in row.cells
        ]
        rows_data.append(cells)

    if not rows_data:
        return ""

    header = rows_data[0]
    if not any(c for c in header):
        return ""

    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows_data[1:]:
        padded = row + [""] * (len(header) - len(row))
        lines.append("| " + " | ".join(padded[:len(header)]) + " |")

    return "\n".join(lines)


# ======================================================================
# 安全校验：文件大小 + zip-bomb 检测
# ======================================================================

def _validate_docx_safe(file_path: str) -> None:
    """在加载 DOCX 前校验文件大小和 ZIP 解压体积，防止 zip-bomb。"""
    file_size = os.path.getsize(file_path)
    if file_size > MAX_DOCX_FILE_SIZE:
        raise ValueError(
            f"DOCX 文件过大（{file_size / 1024 / 1024:.1f} MB），"
            f"上限 {MAX_DOCX_FILE_SIZE // 1024 // 1024} MB"
        )

    if not zipfile.is_zipfile(file_path):
        kind = detect_word_payload_kind(file_path)
        if kind == "ole":
            raise ValueError(
                "检测到旧版 Word（.doc），但自动转换为 .docx 失败。"
                "请用 Word / WPS 另存为 .docx 后再导入。"
            )
        if kind in ("rtf", "html"):
            raise ValueError(
                f"检测到 {kind.upper()} 格式的 Word 文件，但自动转换为 .docx 失败。"
                "请用 Word / WPS 另存为 .docx 后再导入。"
            )
        raise ValueError(UNRECOGNIZED_WORD_PAYLOAD_MSG)

    with zipfile.ZipFile(file_path, "r") as zf:
        total_uncompressed = sum(info.file_size for info in zf.infolist())
        if total_uncompressed > MAX_DOCX_UNCOMPRESSED_SIZE:
            raise ValueError(
                f"DOCX 解压后体积过大（{total_uncompressed / 1024 / 1024:.1f} MB），"
                f"上限 {MAX_DOCX_UNCOMPRESSED_SIZE // 1024 // 1024} MB"
            )
        if file_size > 0 and total_uncompressed / file_size > MAX_DOCX_COMPRESSION_RATIO:
            raise ValueError(
                f"DOCX 压缩比异常（{total_uncompressed / file_size:.0f}:1），"
                f"疑似 zip-bomb，上限 {MAX_DOCX_COMPRESSION_RATIO}:1"
            )

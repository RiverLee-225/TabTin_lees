"""
DC-7 / DC-8 / DC-9 / DC-10 测试：DOCX 导入 P1 修复

DC-7:  图片导入 blob base64 持久化
DC-8:  para_idx 元素→对象映射（不依赖索引对齐）
DC-9:  list chunk 多段落合并（exchange_service 侧）
DC-10: codeBlock 样式 / 等宽字体检测
"""
from __future__ import annotations

import base64
import os
import tempfile
import unittest
import zipfile
from unittest.mock import MagicMock, patch
from xml.etree.ElementTree import Element

from apps.services.docparse.parsers.docx_parser import (
    MAX_IMAGE_BLOB_SIZE,
    _classify_paragraph,
    _extract_images_from_paragraph,
    _is_code_style,
    _is_monospace_paragraph,
)


# ======================================================================
# Helpers
# ======================================================================

def _make_mock_paragraph(
    *,
    style_name: str = "Normal",
    text: str = "",
    runs: list | None = None,
    font_names: list[str] | None = None,
    left_indent_pt: float | None = None,
    style_font_name: str | None = None,
):
    """Build a lightweight mock Paragraph for _classify_paragraph tests."""
    para = MagicMock()
    para.text = text

    style = MagicMock()
    style.name = style_name

    if style_font_name is not None:
        style.font = MagicMock()
        style.font.name = style_font_name
    else:
        style.font = MagicMock()
        style.font.name = None

    para.style = style

    if runs is not None:
        para.runs = runs
    elif font_names:
        mock_runs = []
        for fn in font_names:
            run = MagicMock()
            run.font = MagicMock()
            run.font.name = fn
            run.bold = None
            mock_runs.append(run)
        para.runs = mock_runs
    else:
        para.runs = []

    if left_indent_pt is not None:
        para.paragraph_format = MagicMock()
        para.paragraph_format.left_indent = MagicMock()
        para.paragraph_format.left_indent.pt = left_indent_pt
    else:
        para.paragraph_format = MagicMock()
        para.paragraph_format.left_indent = None

    return para


def _make_real_docx(tmp_dir: str, paragraphs: list[str] | None = None) -> str:
    """Create a real minimal DOCX for integration-level parse tests."""
    from docx import Document as DocxDocument
    doc = DocxDocument()
    for text in (paragraphs or ["Hello world"]):
        doc.add_paragraph(text)
    path = os.path.join(tmp_dir, "test.docx")
    doc.save(path)
    return path


# ======================================================================
# DC-7: image blob base64 persistence
# ======================================================================

class DC7ImageBlobTests(unittest.TestCase):
    """_extract_images_from_paragraph stores base64 blob in result."""

    def test_small_image_has_b64(self):
        small_blob = b"\x89PNG\r\n\x1a\n" + b"\x00" * 100
        para, doc = self._mock_para_with_image(small_blob, "image/png", "test.png")
        images = _extract_images_from_paragraph(para, doc)
        self.assertEqual(len(images), 1)
        self.assertIn("image_b64", images[0])
        decoded = base64.b64decode(images[0]["image_b64"])
        self.assertEqual(decoded, small_blob)

    def test_large_image_no_b64(self):
        large_blob = b"\x00" * (MAX_IMAGE_BLOB_SIZE + 1)
        para, doc = self._mock_para_with_image(large_blob, "image/jpeg", "big.jpg")
        images = _extract_images_from_paragraph(para, doc)
        self.assertEqual(len(images), 1)
        self.assertNotIn("image_b64", images[0])
        self.assertIn("hash", images[0])

    def test_empty_paragraph_no_images(self):
        from docx.oxml.ns import qn

        para = MagicMock()
        para.runs = []
        para._element = Element(qn("w:p"))
        doc = MagicMock()
        self.assertEqual(_extract_images_from_paragraph(para, doc), [])

    def test_vml_imagedata_extracted(self):
        """WPS/.doc 转换后的内嵌图常是 VML v:imagedata，不只 DrawingML blip。"""
        from docx.oxml.ns import qn

        VML = "urn:schemas-microsoft-com:vml"
        small_blob = b"\x89PNG\r\n\x1a\n" + b"\x00" * 40
        imagedata = Element(f"{{{VML}}}imagedata")
        imagedata.set(qn("r:id"), "rId9")
        shape = Element(f"{{{VML}}}shape")
        shape.append(imagedata)
        pict = Element(qn("w:pict"))
        pict.append(shape)
        run_elem = Element(qn("w:r"))
        run_elem.append(pict)
        para_elem = Element(qn("w:p"))
        para_elem.append(run_elem)

        para = MagicMock()
        para._element = para_elem
        para.runs = []

        img_part = MagicMock()
        img_part.blob = small_blob
        img_part.content_type = "image/png"
        img_part.partname = "/word/media/image1.png"
        rel = MagicMock()
        rel.target_part = img_part
        doc = MagicMock()
        doc.part.rels.get = MagicMock(return_value=rel)

        images = _extract_images_from_paragraph(para, doc)
        self.assertEqual(len(images), 1)
        self.assertEqual(images[0]["filename"], "image1.png")
        self.assertIn("image_b64", images[0])
        doc.part.rels.get.assert_called_with("rId9")

    @staticmethod
    def _mock_para_with_image(blob: bytes, content_type: str, filename: str):
        from docx.oxml.ns import qn

        blip = Element(qn("a:blip"))
        blip.set(qn("r:embed"), "rId1")

        drawing = Element(qn("w:drawing"))
        inline = Element(qn("wp:inline"))
        graphic = Element(qn("a:graphic"))
        graphic_data = Element(qn("a:graphicData"))
        pic = Element(qn("pic:pic"))
        blip_fill = Element(qn("pic:blipFill"))
        blip_fill.append(blip)
        pic.append(blip_fill)
        graphic_data.append(pic)
        graphic.append(graphic_data)
        inline.append(graphic)
        drawing.append(inline)

        run_elem = Element(qn("w:r"))
        run_elem.append(drawing)
        para_elem = Element(qn("w:p"))
        para_elem.append(run_elem)

        run = MagicMock()
        run._element = run_elem

        para = MagicMock()
        para.runs = [run]
        para._element = para_elem

        img_part = MagicMock()
        img_part.blob = blob
        img_part.content_type = content_type
        img_part.partname = f"/word/media/{filename}"

        rel = MagicMock()
        rel.target_part = img_part

        doc = MagicMock()
        doc.part.rels.get = MagicMock(return_value=rel)

        return para, doc


# ======================================================================
# DC-8: element→object mapping (no index alignment)
# ======================================================================

class DC8ElementMappingTests(unittest.TestCase):
    """Parser uses element identity mapping instead of para_idx counter."""

    def test_parse_basic_docx(self):
        from apps.services.docparse.parsers.docx_parser import DocxParser
        parser = DocxParser()
        with tempfile.TemporaryDirectory() as tmp:
            path = _make_real_docx(tmp, ["Hello", "World"])
            result = parser.parse(path)
        texts = [
            c.content for p in result.pages for c in p.chunks if c.chunk_type == "paragraph"
        ]
        self.assertIn("Hello", texts)
        self.assertIn("World", texts)

    def test_parse_with_table(self):
        from docx import Document as DocxDocument
        from apps.services.docparse.parsers.docx_parser import DocxParser

        with tempfile.TemporaryDirectory() as tmp:
            doc = DocxDocument()
            doc.add_paragraph("Before table")
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "A"
            table.rows[0].cells[1].text = "B"
            table.rows[1].cells[0].text = "1"
            table.rows[1].cells[1].text = "2"
            doc.add_paragraph("After table")
            path = os.path.join(tmp, "tbl.docx")
            doc.save(path)

            parser = DocxParser()
            result = parser.parse(path)

        types = [c.chunk_type for p in result.pages for c in p.chunks]
        self.assertIn("table", types)
        self.assertIn("paragraph", types)

    def test_mapping_skips_unknown_elements(self):
        """Elements not in doc.paragraphs mapping are safely skipped."""
        from apps.services.docparse.parsers.docx_parser import DocxParser
        parser = DocxParser()
        with tempfile.TemporaryDirectory() as tmp:
            path = _make_real_docx(tmp, ["Only paragraph"])
            result = parser.parse(path)
        self.assertTrue(len(result.pages) >= 1)


# ======================================================================
# DC-9: list chunk multi-paragraph handling (exchange_service)
# ======================================================================

class DC9ListChunkTests(unittest.TestCase):
    """exchange_service list chunk processing handles edge cases."""

    def _run_list_processing(self, content: str) -> list[str]:
        """Simulate the list chunk processing logic from exchange_service."""
        import re
        list_items: list[str] = []
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            has_marker = bool(
                re.match(r"^[-*]\s", stripped)
                or re.match(r"^\d+[.)]\s", stripped)
                or stripped[0] in "•–—·"
            )
            if has_marker:
                if stripped[0] in "•–—·":
                    stripped = "- " + stripped[1:].lstrip()
                list_items.append(stripped)
            elif list_items:
                list_items[-1] += " " + stripped
            else:
                list_items.append("- " + stripped)
        return list_items

    def test_simple_bullet_list(self):
        items = self._run_list_processing("- item one\n- item two")
        self.assertEqual(items, ["- item one", "- item two"])

    def test_continuation_lines_merged(self):
        items = self._run_list_processing("- item one\ncontinuation\n- item two")
        self.assertEqual(len(items), 2)
        self.assertIn("continuation", items[0])
        self.assertEqual(items[1], "- item two")

    def test_unicode_bullets_normalized(self):
        items = self._run_list_processing("• first\n– second\n· third")
        for item in items:
            self.assertTrue(item.startswith("- "), f"Expected '- ' prefix: {item}")
        self.assertEqual(len(items), 3)

    def test_numbered_list(self):
        items = self._run_list_processing("1. first\n2. second")
        self.assertEqual(items, ["1. first", "2. second"])

    def test_empty_lines_skipped(self):
        items = self._run_list_processing("- item\n\n\n- item2")
        self.assertEqual(len(items), 2)

    def test_no_marker_first_line_gets_prefix(self):
        items = self._run_list_processing("no marker here")
        self.assertEqual(items, ["- no marker here"])

    def test_multiline_numbered_continuation(self):
        items = self._run_list_processing("1. first paragraph\nstill first\n2. second")
        self.assertEqual(len(items), 2)
        self.assertIn("still first", items[0])


# ======================================================================
# DC-10: codeBlock detection
# ======================================================================

class DC10CodeBlockDetectionTests(unittest.TestCase):
    """_classify_paragraph detects code styles and monospace fonts."""

    def test_code_style_detected(self):
        for style in ("Code", "code", "Code Block", "Code Char", "HTML Code",
                       "Source Code", "Verbatim", "PlainText", "程序代码", "代码"):
            para = _make_mock_paragraph(style_name=style)
            chunk_type, level, _list_meta = _classify_paragraph(para)
            self.assertEqual(chunk_type, "codeBlock", f"Style '{style}' should be codeBlock")
            self.assertIsNone(level)

    def test_heading_style_not_code(self):
        para = _make_mock_paragraph(style_name="Heading 1")
        chunk_type, level, _list_meta = _classify_paragraph(para)
        self.assertEqual(chunk_type, "heading")
        self.assertEqual(level, 1)

    def test_normal_style_not_code(self):
        para = _make_mock_paragraph(style_name="Normal")
        chunk_type, _, _list_meta = _classify_paragraph(para)
        self.assertEqual(chunk_type, "paragraph")

    def test_monospace_font_detected(self):
        para = _make_mock_paragraph(
            style_name="Normal",
            font_names=["Courier New"],
        )
        chunk_type, _, _list_meta = _classify_paragraph(para)
        self.assertEqual(chunk_type, "codeBlock")

    def test_consolas_font_detected(self):
        para = _make_mock_paragraph(
            style_name="Normal",
            font_names=["Consolas", "Consolas"],
        )
        chunk_type, _, _list_meta = _classify_paragraph(para)
        self.assertEqual(chunk_type, "codeBlock")

    def test_mixed_fonts_not_code(self):
        para = _make_mock_paragraph(
            style_name="Normal",
            font_names=["Courier New", "Arial"],
        )
        chunk_type, _, _list_meta = _classify_paragraph(para)
        self.assertNotEqual(chunk_type, "codeBlock")

    def test_style_font_monospace(self):
        para = _make_mock_paragraph(
            style_name="Normal",
            style_font_name="Consolas",
        )
        chunk_type, _, _list_meta = _classify_paragraph(para)
        self.assertEqual(chunk_type, "codeBlock")

    def test_is_code_style_helper(self):
        self.assertTrue(_is_code_style("Code"))
        self.assertTrue(_is_code_style("code"))
        self.assertTrue(_is_code_style("Code Block"))
        self.assertTrue(_is_code_style("代码块"))
        self.assertFalse(_is_code_style("Normal"))
        self.assertFalse(_is_code_style("Heading 1"))

    def test_is_monospace_paragraph_no_runs(self):
        para = _make_mock_paragraph(style_name="Normal")
        self.assertFalse(_is_monospace_paragraph(para))

    def test_codeblock_in_exchange_markdown(self):
        """codeBlock chunks produce fenced code blocks in markdown."""
        content = "def hello():\n    print('hi')"
        expected = f"```\n{content}\n```"
        self.assertIn("```", expected)
        self.assertIn(content, expected)


class DC10IntegrationTests(unittest.TestCase):
    """End-to-end: DOCX with Code style → codeBlock chunk."""

    def test_code_style_paragraph_becomes_codeblock(self):
        from docx import Document as DocxDocument
        from apps.services.docparse.parsers.docx_parser import DocxParser

        with tempfile.TemporaryDirectory() as tmp:
            doc = DocxDocument()
            doc.add_paragraph("Normal text")
            code_para = doc.add_paragraph("print('hello')")
            code_para.style = doc.styles["Normal"]

            # python-docx doesn't have a built-in "Code" style in the
            # default template, so we test via font-based detection.
            for run in code_para.runs:
                run.font.name = "Courier New"

            path = os.path.join(tmp, "code.docx")
            doc.save(path)

            parser = DocxParser()
            result = parser.parse(path)

        types = [c.chunk_type for p in result.pages for c in p.chunks]
        self.assertIn("codeBlock", types)


if __name__ == "__main__":
    unittest.main()

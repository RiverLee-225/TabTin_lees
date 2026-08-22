"""DC-12：docx 导入产出真正的 bulletList/orderedList（Refs ）。

覆盖三层：
  1. docx_parser 对 numbering.xml（numId/ilvl）的解析——区分 bullet/ordered、
     保留层级、样式级 numPr 回退、启发式兜底的符号裁剪。
  2. DocxParser 端到端解析真实 .docx，验证 chunk_type="list" 携带的
     list_kind/list_level/list_num_id 元数据。
  3. service._build_tabdoc_import_draft 把携带这套元数据的 chunk 序列组装成
     嵌套 bulletList/orderedList/listItem pm_json 树，并在非 docx 列表
     （缺这套元数据）时退回原有 markdown 中转路径。
"""

from __future__ import annotations

import os
import tempfile
import unittest
from types import SimpleNamespace

from django.test import TestCase
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement

from apps.services.docparse.models import DocumentChunk, DocumentPage, ParsedDocument
from apps.services.docparse.parsers.docx_parser import (
    DocxParser,
    _classify_paragraph,
    _heuristic_list_kind,
    _load_numbering_definitions,
    _paragraph_num_pr,
    _resolve_list_info,
    _strip_list_marker_from_runs,
)
from apps.services.docparse.service import (
    _build_tabdoc_import_draft,
    _docx_list_item_pm_content,
    _docx_list_meta,
)
from apps.services.docparse.tests_import_job_execution import _file_record


def _add_custom_numbering(doc, levels: dict[int, str]) -> str:
    """在 numbering.xml 里插入一个自定义 abstractNum + num，
    levels 形如 ``{ilvl: numFmt}``（如 ``{0: "bullet", 1: "decimal"}``）。
    返回新建 num 的 numId（字符串）。
    """
    numbering_elm = doc.part.numbering_part.element

    existing_abstract_ids = [
        int(a.get(qn("w:abstractNumId")))
        for a in numbering_elm.findall(qn("w:abstractNum"))
        if a.get(qn("w:abstractNumId")) is not None
    ]
    abstract_id = max(existing_abstract_ids, default=-1) + 1

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
    for ilvl, num_fmt in levels.items():
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl.append(fmt)
        abstract_num.append(lvl)
    numbering_elm.insert(0, abstract_num)

    existing_num_ids = [
        int(n.get(qn("w:numId")))
        for n in numbering_elm.findall(qn("w:num"))
        if n.get(qn("w:numId")) is not None
    ]
    num_id = max(existing_num_ids, default=0) + 1

    num_elem = OxmlElement("w:num")
    num_elem.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num_elem.append(abstract_ref)
    numbering_elm.append(num_elem)

    return str(num_id)


def _set_paragraph_numbering(para, num_id: str, ilvl: int) -> None:
    """给段落直接挂 <w:pPr><w:numPr>（绕开样式，精确控制 numId/ilvl）。"""
    pPr = para._element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_elem = OxmlElement("w:ilvl")
    ilvl_elem.set(qn("w:val"), str(ilvl))
    num_pr.append(ilvl_elem)
    num_id_elem = OxmlElement("w:numId")
    num_id_elem.set(qn("w:val"), num_id)
    num_pr.append(num_id_elem)
    pPr.append(num_pr)


class NumberingDefinitionUnitTests(unittest.TestCase):
    """针对 numbering.xml 解析 + numPr 判定的纯函数单测。"""

    def _build_doc_with_numbering(self, levels: dict[int, str]):
        from docx import Document as DocxDocument

        doc = DocxDocument()
        num_id = _add_custom_numbering(doc, levels)
        return doc, num_id

    def test_load_numbering_definitions_distinguishes_bullet_and_ordered(self):
        doc, num_id = self._build_doc_with_numbering({0: "bullet", 1: "decimal"})
        defs = _load_numbering_definitions(doc)
        self.assertEqual(defs[num_id][0]["kind"], "bullet")
        self.assertEqual(defs[num_id][1]["kind"], "ordered")

    def test_load_numbering_definitions_treats_none_fmt_as_non_list(self):
        doc, num_id = self._build_doc_with_numbering({0: "none"})
        defs = _load_numbering_definitions(doc)
        self.assertEqual(defs[num_id][0]["kind"], "none")

    def test_resolve_list_info_returns_none_without_numpr(self):
        doc, num_id = self._build_doc_with_numbering({0: "bullet"})
        para = doc.add_paragraph("plain paragraph, no numPr")
        defs = _load_numbering_definitions(doc)
        self.assertIsNone(_resolve_list_info(defs, para))

    def test_resolve_list_info_returns_none_for_numfmt_none(self):
        doc, num_id = self._build_doc_with_numbering({0: "none"})
        para = doc.add_paragraph("removed numbering placeholder")
        _set_paragraph_numbering(para, num_id, 0)
        defs = _load_numbering_definitions(doc)
        self.assertIsNone(_resolve_list_info(defs, para))

    def test_resolve_list_info_reports_kind_level_and_start(self):
        doc, num_id = self._build_doc_with_numbering({0: "bullet", 1: "decimal"})
        para = doc.add_paragraph("nested ordered item")
        _set_paragraph_numbering(para, num_id, 1)
        defs = _load_numbering_definitions(doc)
        info = _resolve_list_info(defs, para)
        self.assertEqual(info["kind"], "ordered")
        self.assertEqual(info["level"], 1)
        self.assertEqual(info["num_id"], num_id)

    def test_paragraph_num_pr_numid_zero_means_no_list(self):
        doc, _num_id = self._build_doc_with_numbering({0: "bullet"})
        para = doc.add_paragraph("explicit numbering removal")
        _set_paragraph_numbering(para, "0", 0)
        self.assertIsNone(_paragraph_num_pr(para))

    def test_paragraph_num_pr_falls_back_to_style_level_numpr(self):
        """Word 内置 "List Bullet"/"List Number" 把 numPr 挂在样式而非段落上。"""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        para = doc.add_paragraph("style level numbering", style="List Bullet")
        result = _paragraph_num_pr(para)
        self.assertIsNotNone(result)
        num_id, ilvl = result
        self.assertEqual(ilvl, 0)
        defs = _load_numbering_definitions(doc)
        self.assertEqual(defs[num_id][0]["kind"], "bullet")


class ClassifyParagraphListTests(unittest.TestCase):
    """_classify_paragraph 对 list_meta 的组装（numPr 优先于启发式兜底）。"""

    def _fake_para(self, text: str, style_name: str = "Normal", left_indent_pt: float | None = None):
        paragraph_format = SimpleNamespace(
            left_indent=SimpleNamespace(pt=left_indent_pt) if left_indent_pt else None
        )
        return SimpleNamespace(
            style=SimpleNamespace(name=style_name),
            text=text,
            runs=[],
            paragraph_format=paragraph_format,
        )

    def test_numpr_list_info_wins_over_heuristic(self):
        para = self._fake_para("Item text", style_name="Normal")
        list_info = {"kind": "ordered", "level": 2, "num_id": "7", "start": 3}
        chunk_type, heading_level, list_meta = _classify_paragraph(para, list_info)
        self.assertEqual(chunk_type, "list")
        self.assertIsNone(heading_level)
        self.assertEqual(list_meta["list_kind"], "ordered")
        self.assertEqual(list_meta["list_level"], 2)
        self.assertEqual(list_meta["list_num_id"], "7")
        self.assertEqual(list_meta["list_start"], 3)

    def test_heuristic_bullet_char_without_numpr(self):
        para = self._fake_para("• bullet via style name", style_name="List Paragraph")
        chunk_type, _level, list_meta = _classify_paragraph(para, None)
        self.assertEqual(chunk_type, "list")
        self.assertEqual(list_meta["list_kind"], "bullet")
        self.assertEqual(list_meta["list_level"], 0)
        self.assertIsNone(list_meta["list_num_id"])

    def test_heuristic_ordered_digit_prefix_via_indent(self):
        para = self._fake_para("1. ordered via indent", style_name="Normal", left_indent_pt=30)
        chunk_type, _level, list_meta = _classify_paragraph(para, None)
        self.assertEqual(chunk_type, "list")
        self.assertEqual(list_meta["list_kind"], "ordered")

    def test_plain_paragraph_without_list_signal(self):
        para = self._fake_para("just a paragraph", style_name="Normal")
        chunk_type, _level, list_meta = _classify_paragraph(para, None)
        self.assertEqual(chunk_type, "paragraph")
        self.assertIsNone(list_meta)

    def test_heuristic_list_kind_helper(self):
        self.assertEqual(_heuristic_list_kind("1. ordered"), "ordered")
        self.assertEqual(_heuristic_list_kind("• bullet"), "bullet")


class StripListMarkerTests(unittest.TestCase):
    def test_strips_bullet_char_from_first_run(self):
        runs_meta = [{"text": "• Item one"}]
        adjusted = _strip_list_marker_from_runs(runs_meta, "• Item one")
        self.assertEqual(adjusted[0]["text"], "Item one")

    def test_strips_ordered_prefix_across_run_boundary(self):
        runs_meta = [{"text": "1."}, {"text": " Item", "bold": True}]
        adjusted = _strip_list_marker_from_runs(runs_meta, "1. Item")
        self.assertEqual([r["text"] for r in adjusted], ["Item"])
        self.assertTrue(adjusted[0]["bold"])

    def test_no_marker_leaves_runs_untouched(self):
        runs_meta = [{"text": "no marker here"}]
        adjusted = _strip_list_marker_from_runs(runs_meta, "no marker here")
        self.assertEqual(adjusted, runs_meta)


class DocxParserListIntegrationTests(unittest.TestCase):
    """用真实 python-docx 文档端到端验证 DocxParser 的列表识别。"""

    def test_numpr_bullet_and_nested_ordered_levels_are_detected(self):
        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc = DocxDocument()
            num_id = _add_custom_numbering(doc, {0: "bullet", 1: "decimal"})

            p1 = doc.add_paragraph()
            run = p1.add_run("Top level bullet")
            run.bold = True
            _set_paragraph_numbering(p1, num_id, 0)

            p2 = doc.add_paragraph("Nested ordered item")
            _set_paragraph_numbering(p2, num_id, 1)

            path = os.path.join(tmp_dir, "numbered-list.docx")
            doc.save(path)

            result = DocxParser()._parse_docx(path)
            chunks = [c for page in result.pages for c in page.chunks]
            list_chunks = [c for c in chunks if c.chunk_type == "list"]
            self.assertEqual(len(list_chunks), 2)

            top, nested = list_chunks
            self.assertEqual(top.metadata["list_kind"], "bullet")
            self.assertEqual(top.metadata["list_level"], 0)
            self.assertEqual(top.metadata["list_num_id"], num_id)
            self.assertTrue(top.metadata["runs"][0]["bold"])

            self.assertEqual(nested.metadata["list_kind"], "ordered")
            self.assertEqual(nested.metadata["list_level"], 1)
            self.assertEqual(nested.metadata["list_num_id"], num_id)

    def test_builtin_list_bullet_style_detected_via_style_level_numpr(self):
        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc = DocxDocument()
            doc.add_paragraph("Built-in bullet style item", style="List Bullet")
            doc.add_paragraph("Built-in number style item", style="List Number")
            path = os.path.join(tmp_dir, "builtin-list-styles.docx")
            doc.save(path)

            result = DocxParser()._parse_docx(path)
            chunks = [c for page in result.pages for c in page.chunks]
            list_chunks = [c for c in chunks if c.chunk_type == "list"]
            self.assertEqual(len(list_chunks), 2)
            self.assertEqual(list_chunks[0].metadata["list_kind"], "bullet")
            self.assertEqual(list_chunks[1].metadata["list_kind"], "ordered")

    def test_heuristic_list_without_numpr_strips_literal_marker_from_runs(self):
        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc = DocxDocument()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = None
            from docx.shared import Pt

            para.paragraph_format.left_indent = Pt(30)
            para.add_run("• literal bullet char in text")
            path = os.path.join(tmp_dir, "heuristic-list.docx")
            doc.save(path)

            result = DocxParser()._parse_docx(path)
            chunks = [c for page in result.pages for c in page.chunks]
            list_chunks = [c for c in chunks if c.chunk_type == "list"]
            self.assertEqual(len(list_chunks), 1)
            chunk = list_chunks[0]
            self.assertIsNone(chunk.metadata.get("list_num_id"))
            runs_meta = chunk.metadata.get("runs") or []
            self.assertTrue(runs_meta)
            self.assertEqual(runs_meta[0]["text"], "literal bullet char in text")


class ImportDraftListAssemblyTests(TestCase):
    """service._build_tabdoc_import_draft 把 docx 列表 chunk 序列组装成嵌套
    bulletList/orderedList/listItem pm_json 树（Refs ）。
    """

    def _make_parsed_document(self) -> ParsedDocument:
        file_record = _file_record("list-fidelity.docx")
        return ParsedDocument.objects.create(
            file_record=file_record,
            status=ParsedDocument.Status.READY,
            title="list fidelity",
            total_pages=1,
            parsed_pages=1,
        )

    def test_nested_bullet_and_ordered_lists_with_dedent(self):
        parsed = self._make_parsed_document()
        page = DocumentPage.objects.create(document=parsed, page_number=1)

        DocumentChunk.objects.create(
            page=page, chunk_type=DocumentChunk.ChunkType.LIST, content="Bullet one",
            sequence=1,
            metadata={
                "list_kind": "bullet", "list_level": 0, "list_num_id": "10",
                "runs": [{"text": "Bullet one", "bold": True}],
            },
        )
        DocumentChunk.objects.create(
            page=page, chunk_type=DocumentChunk.ChunkType.LIST, content="Bullet two",
            sequence=2,
            metadata={"list_kind": "bullet", "list_level": 0, "list_num_id": "10"},
        )
        DocumentChunk.objects.create(
            page=page, chunk_type=DocumentChunk.ChunkType.LIST, content="Nested ordered a",
            sequence=3,
            metadata={"list_kind": "ordered", "list_level": 1, "list_num_id": "11", "list_start": 1},
        )
        DocumentChunk.objects.create(
            page=page, chunk_type=DocumentChunk.ChunkType.LIST, content="Nested ordered b",
            sequence=4,
            metadata={"list_kind": "ordered", "list_level": 1, "list_num_id": "11", "list_start": 1},
        )
        DocumentChunk.objects.create(
            page=page, chunk_type=DocumentChunk.ChunkType.LIST, content="Bullet three (dedent back)",
            sequence=5,
            metadata={"list_kind": "bullet", "list_level": 0, "list_num_id": "10"},
        )

        draft = _build_tabdoc_import_draft(parsed)
        content = draft["pm_json"]["content"]

        self.assertEqual(len(content), 1)
        top_list = content[0]
        self.assertEqual(top_list["type"], "bulletList")
        self.assertEqual(len(top_list["content"]), 3)

        item1, item2, item3 = top_list["content"]
        self.assertEqual(item1["type"], "listItem")
        self.assertEqual(item1["content"][0]["content"][0]["text"], "Bullet one")
        self.assertEqual(item1["content"][0]["content"][0]["marks"], [{"type": "bold"}])

        # Item 2 里嵌套的 ordered 列表
        nested_children = [c for c in item2["content"] if c["type"] == "orderedList"]
        self.assertEqual(len(nested_children), 1)
        nested_list = nested_children[0]
        self.assertEqual(nested_list["attrs"]["start"], 1)
        self.assertEqual(len(nested_list["content"]), 2)
        self.assertEqual(
            nested_list["content"][0]["content"][0]["content"][0]["text"],
            "Nested ordered a",
        )

        self.assertEqual(item3["content"][0]["content"][0]["text"], "Bullet three (dedent back)")

    def test_paragraph_between_same_numid_lists_starts_a_new_list_block(self):
        parsed = self._make_parsed_document()
        page = DocumentPage.objects.create(document=parsed, page_number=1)

        DocumentChunk.objects.create(
            page=page, chunk_type=DocumentChunk.ChunkType.LIST, content="First list item",
            sequence=1,
            metadata={"list_kind": "bullet", "list_level": 0, "list_num_id": "20"},
        )
        DocumentChunk.objects.create(
            page=page, chunk_type=DocumentChunk.ChunkType.PARAGRAPH, content="An interrupting paragraph",
            sequence=2,
            metadata={"runs": [{"text": "An interrupting paragraph"}]},
        )
        DocumentChunk.objects.create(
            page=page, chunk_type=DocumentChunk.ChunkType.LIST, content="Second list item",
            sequence=3,
            metadata={"list_kind": "bullet", "list_level": 0, "list_num_id": "20"},
        )

        draft = _build_tabdoc_import_draft(parsed)
        content = draft["pm_json"]["content"]

        list_nodes = [n for n in content if n["type"] == "bulletList"]
        self.assertEqual(len(list_nodes), 2)
        self.assertEqual(len(list_nodes[0]["content"]), 1)
        self.assertEqual(len(list_nodes[1]["content"]), 1)

    def test_list_chunk_without_docx_metadata_falls_back_to_markdown_path(self):
        """PDF/纯文本解析器产出的 "list" chunk 没有 list_kind 元数据，
        应继续走原有 markdown_to_pm_json 中转路径，不应崩溃或被吞掉。
        """
        parsed = self._make_parsed_document()
        page = DocumentPage.objects.create(document=parsed, page_number=1)
        DocumentChunk.objects.create(
            page=page, chunk_type=DocumentChunk.ChunkType.LIST, content="- markdown bullet already",
            sequence=1,
            metadata={"source": "text_layer"},
        )

        draft = _build_tabdoc_import_draft(parsed)
        content = draft["pm_json"]["content"]
        self.assertTrue(content)
        self.assertEqual(content[0]["type"], "bulletList")

    def test_docx_list_meta_and_item_content_helpers(self):
        chunk = SimpleNamespace(
            chunk_type=DocumentChunk.ChunkType.LIST,
            content="Styled item",
            metadata={
                "list_kind": "ordered", "list_level": 0, "list_num_id": "5", "list_start": 4,
                "runs": [{"text": "Styled item", "italic": True}],
            },
        )
        meta = _docx_list_meta(chunk)
        self.assertEqual(meta, {"kind": "ordered", "level": 0, "numid": "5", "start": 4})

        item_content = _docx_list_item_pm_content(chunk)
        self.assertEqual(item_content[0]["type"], "paragraph")
        self.assertEqual(item_content[0]["content"][0]["marks"], [{"type": "italic"}])

    def test_docx_list_meta_returns_none_for_non_docx_list_chunk(self):
        chunk = SimpleNamespace(
            chunk_type=DocumentChunk.ChunkType.LIST,
            content="plain list",
            metadata={"source": "text_layer"},
        )
        self.assertIsNone(_docx_list_meta(chunk))


if __name__ == "__main__":
    unittest.main()

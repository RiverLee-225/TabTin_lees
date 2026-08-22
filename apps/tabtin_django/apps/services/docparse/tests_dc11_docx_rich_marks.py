"""DC-11：docx run 级样式保真 → pmJson marks。

Refs ：补充标题对齐样式继承、下划线样式继承、主题色近似兜底的测试。
"""

from __future__ import annotations

import os
import tempfile
import unittest
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

from django.test import TestCase

from apps.services.docparse.models import DocumentChunk, DocumentPage, ParsedDocument
from apps.services.docparse.parsers.docx_parser import (
    DocxParser,
    _extract_run_marks,
    _paragraph_align,
    _run_color_hex,
    _run_marks,
    _run_underline,
)
from apps.services.docparse.service import (
    _build_tabdoc_import_draft,
    _docx_rich_pm_node,
    _pm_text_nodes_from_docx_runs,
)
from apps.services.docparse.tests_import_job_execution import _file_record
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX


class FakeRGB:
    def __init__(self, value: str):
        self._value = value

    def __str__(self) -> str:
        return self._value


def _fake_run(
    text: str,
    *,
    bold=None,
    italic=None,
    underline=None,
    strike=False,
    superscript=False,
    subscript=False,
    color_hex: str | None = None,
    highlight=None,
    style_name: str | None = None,
):
    font = SimpleNamespace(
        strike=strike,
        double_strike=False,
        superscript=superscript,
        subscript=subscript,
        highlight_color=highlight,
        color=SimpleNamespace(
            type=True if color_hex else None,
            rgb=FakeRGB(color_hex) if color_hex else None,
        ),
        name=None,
    )
    style = SimpleNamespace(name=style_name) if style_name else None
    element = MagicMock()
    element.find.return_value = None
    return SimpleNamespace(
        text=text,
        bold=bold,
        italic=italic,
        underline=underline,
        font=font,
        style=style,
        _element=element,
    )


class RunMarksUnitTests(unittest.TestCase):
    def test_run_marks_capture_inline_styles(self):
        run = _fake_run(
            "red bold",
            bold=True,
            italic=True,
            underline=True,
            strike=True,
            superscript=True,
            color_hex="FF0000",
            highlight=WD_COLOR_INDEX.YELLOW,
        )
        marks = _run_marks(run)
        self.assertTrue(marks["bold"])
        self.assertTrue(marks["italic"])
        self.assertTrue(marks["underline"])
        self.assertTrue(marks["strike"])
        self.assertTrue(marks["superscript"])
        self.assertEqual(marks["color"], "#FF0000")
        self.assertEqual(marks["highlight"], "#fef9c3")

    def test_char_style_fallback_for_emphasis(self):
        run = _fake_run("subtle", style_name="Subtle Emphasis")
        marks = _run_marks(run)
        self.assertTrue(marks["italic"])
        self.assertEqual(marks["color"], "#8C8C8C")

    def test_paragraph_align_right(self):
        para = SimpleNamespace(alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(_paragraph_align(para), "right")

    def test_theme_color_approximation_when_no_explicit_rgb(self):
        """Refs ：无显式 RGB 快照、仅引用主题色时按默认主题近似取色。"""
        theme_color = SimpleNamespace(xml_value="accent1")
        run = _fake_run("themed")
        run.font.color.type = True
        run.font.color.rgb = None
        run.font.color.theme_color = theme_color
        self.assertEqual(_run_color_hex(run), "#4472C4")

    def test_explicit_rgb_wins_over_theme_color(self):
        run = _fake_run("explicit", color_hex="112233")
        run.font.color.theme_color = SimpleNamespace(xml_value="accent1")
        self.assertEqual(_run_color_hex(run), "#112233")

    def test_inherits_color_from_paragraph_heading_style(self):
        """Refs ：标题蓝字写在 Heading 样式上、run 无 rPr 时要继承。"""
        para_style = SimpleNamespace(
            font=SimpleNamespace(
                color=SimpleNamespace(type=True, rgb=FakeRGB("2E75B5"), theme_color=None),
            ),
            base_style=None,
        )
        run = _fake_run("Structural Elements")
        run.font.color.type = None
        run.font.color.rgb = None
        self.assertEqual(_run_color_hex(run, para_style), "#2E75B5")
        marks = _run_marks(run, para_style)
        self.assertEqual(marks.get("color"), "#2E75B5")


class ParagraphAlignStyleInheritanceTests(unittest.TestCase):
    """Refs ：para.alignment 为空时回退 style.paragraph_format.alignment。"""

    def test_falls_back_to_style_alignment_when_direct_format_unset(self):
        style = SimpleNamespace(
            paragraph_format=SimpleNamespace(alignment=WD_ALIGN_PARAGRAPH.CENTER),
            base_style=None,
        )
        para = SimpleNamespace(alignment=None, style=style)
        self.assertEqual(_paragraph_align(para), "center")

    def test_walks_base_style_chain_when_own_style_unset(self):
        grandparent_style = SimpleNamespace(
            paragraph_format=SimpleNamespace(alignment=WD_ALIGN_PARAGRAPH.RIGHT),
            base_style=None,
        )
        style = SimpleNamespace(
            paragraph_format=SimpleNamespace(alignment=None),
            base_style=grandparent_style,
        )
        para = SimpleNamespace(alignment=None, style=style)
        self.assertEqual(_paragraph_align(para), "right")

    def test_direct_paragraph_format_takes_precedence_over_style(self):
        style = SimpleNamespace(
            paragraph_format=SimpleNamespace(alignment=WD_ALIGN_PARAGRAPH.CENTER),
            base_style=None,
        )
        para = SimpleNamespace(alignment=WD_ALIGN_PARAGRAPH.RIGHT, style=style)
        self.assertEqual(_paragraph_align(para), "right")

    def test_no_alignment_anywhere_returns_none(self):
        style = SimpleNamespace(
            paragraph_format=SimpleNamespace(alignment=None),
            base_style=None,
        )
        para = SimpleNamespace(alignment=None, style=style)
        self.assertIsNone(_paragraph_align(para))


class RunUnderlineInheritanceTests(unittest.TestCase):
    """Refs ：run 未显式设置 underline 时按字符样式/段落样式链回退。"""

    def test_direct_underline_wins(self):
        run = SimpleNamespace(underline=True, style=None)
        self.assertTrue(_run_underline(run))

    def test_direct_false_overrides_inherited_style(self):
        char_style = SimpleNamespace(font=SimpleNamespace(underline=True), base_style=None)
        run = SimpleNamespace(underline=False, style=char_style)
        self.assertFalse(_run_underline(run))

    def test_inherits_from_character_style(self):
        char_style = SimpleNamespace(font=SimpleNamespace(underline=True), base_style=None)
        run = SimpleNamespace(underline=None, style=char_style)
        self.assertTrue(_run_underline(run))

    def test_character_style_chain_via_base_style(self):
        grandparent = SimpleNamespace(font=SimpleNamespace(underline=True), base_style=None)
        char_style = SimpleNamespace(font=SimpleNamespace(underline=None), base_style=grandparent)
        run = SimpleNamespace(underline=None, style=char_style)
        self.assertTrue(_run_underline(run))

    def test_falls_back_to_paragraph_style_when_char_style_unset(self):
        char_style = SimpleNamespace(font=SimpleNamespace(underline=None), base_style=None)
        para_style = SimpleNamespace(font=SimpleNamespace(underline=True), base_style=None)
        run = SimpleNamespace(underline=None, style=char_style)
        self.assertTrue(_run_underline(run, para_style))

    def test_no_underline_anywhere_returns_false(self):
        char_style = SimpleNamespace(font=SimpleNamespace(underline=None), base_style=None)
        para_style = SimpleNamespace(font=SimpleNamespace(underline=None), base_style=None)
        run = SimpleNamespace(underline=None, style=char_style)
        self.assertFalse(_run_underline(run, para_style))

    def test_run_marks_reports_inherited_underline(self):
        char_style = SimpleNamespace(font=SimpleNamespace(underline=True), base_style=None)
        run = _fake_run("linked", style_name=None)
        run.style = char_style
        marks = _run_marks(run)
        self.assertTrue(marks.get("underline"))


class PmJsonFromRunsTests(unittest.TestCase):
    def test_pm_text_nodes_map_marks(self):
        nodes = _pm_text_nodes_from_docx_runs([
            {
                "text": "Hello",
                "bold": True,
                "color": "#E00000",
                "highlight": "#fef9c3",
                "link": "https://example.com",
            },
            {"text": " world", "italic": True},
        ])
        self.assertEqual(len(nodes), 2)
        marks = {m["type"]: m for m in nodes[0]["marks"]}
        self.assertIn("bold", marks)
        self.assertEqual(marks["textStyle"]["attrs"]["color"], "#E00000")
        self.assertEqual(marks["highlight"]["attrs"]["color"], "#fef9c3")
        self.assertEqual(marks["link"]["attrs"]["href"], "https://example.com")
        self.assertEqual(nodes[1]["marks"], [{"type": "italic"}])

    def test_rich_pm_node_for_paragraph_with_align(self):
        chunk = SimpleNamespace(
            chunk_type=DocumentChunk.ChunkType.PARAGRAPH,
            heading_level=None,
            metadata={
                "align": "right",
                "runs": [{"text": "Right aligned", "bold": True}],
            },
        )
        node = _docx_rich_pm_node(chunk)
        self.assertEqual(node["type"], "paragraph")
        self.assertEqual(node["attrs"]["textAlign"], "right")
        self.assertEqual(node["content"][0]["text"], "Right aligned")
        self.assertEqual(node["content"][0]["marks"], [{"type": "bold"}])


class ImportDraftRichDocxTests(TestCase):
    @patch("apps.services.docparse.service._build_import_image_markdown")
    def test_import_draft_preserves_docx_run_marks_in_pm_json(self, _mock_img):
        file_record = _file_record("styles.docx")
        parsed = ParsedDocument.objects.create(
            file_record=file_record,
            status=ParsedDocument.Status.READY,
            title="Demonstration of DOCX support in calibre",
            total_pages=1,
            parsed_pages=1,
        )
        page = DocumentPage.objects.create(document=parsed, page_number=1)
        DocumentChunk.objects.create(
            page=page,
            chunk_type=DocumentChunk.ChunkType.HEADING,
            content="Demonstration of DOCX support in calibre",
            sequence=1,
            heading_level=1,
            metadata={
                "runs": [{"text": "Demonstration of DOCX support in calibre", "bold": True}],
            },
        )
        DocumentChunk.objects.create(
            page=page,
            chunk_type=DocumentChunk.ChunkType.PARAGRAPH,
            content="This is bold, italic, and red.",
            sequence=2,
            metadata={
                "align": "right",
                "runs": [
                    {"text": "This is "},
                    {"text": "bold", "bold": True},
                    {"text": ", "},
                    {"text": "italic", "italic": True},
                    {"text": ", and "},
                    {"text": "red", "color": "#FF0000", "highlight": "#fef9c3"},
                    {"text": "."},
                ],
            },
        )

        draft = _build_tabdoc_import_draft(parsed)
        content = draft["pm_json"]["content"]
        self.assertEqual(content[0]["type"], "heading")
        self.assertEqual(content[1]["type"], "paragraph")
        self.assertEqual(content[1]["attrs"]["textAlign"], "right")

        mark_types: set[str] = set()
        colors: list[str] = []
        highlights: list[str] = []

        def walk(node):
            if isinstance(node, dict):
                for mark in node.get("marks") or []:
                    mark_types.add(mark.get("type"))
                    if mark.get("type") == "textStyle":
                        colors.append(mark.get("attrs", {}).get("color"))
                    if mark.get("type") == "highlight":
                        highlights.append(mark.get("attrs", {}).get("color"))
                for child in node.get("content") or []:
                    walk(child)
            elif isinstance(node, list):
                for child in node:
                    walk(child)

        walk(content)
        self.assertIn("bold", mark_types)
        self.assertIn("italic", mark_types)
        self.assertIn("textStyle", mark_types)
        self.assertIn("highlight", mark_types)
        self.assertIn("#FF0000", colors)
        self.assertIn("#fef9c3", highlights)


class DocxParserStyleInheritanceIntegrationTests(unittest.TestCase):
    """Refs ：用真实 python-docx 文档端到端验证对齐/下划线的样式继承。"""

    def test_heading_alignment_defined_only_at_style_level_is_preserved(self):
        """标题居中常见写法：只改样式的 paragraph_format.alignment，不逐段设置直接格式。"""
        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc = DocxDocument()
            heading_style = doc.styles["Heading 1"]
            heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("Centered Heading", style="Heading 1")
            path = os.path.join(tmp_dir, "styled-heading.docx")
            doc.save(path)

            result = DocxParser()._parse_docx(path)
            chunks = [c for page in result.pages for c in page.chunks]
            heading_chunk = next(c for c in chunks if c.chunk_type == "heading")
            self.assertEqual(heading_chunk.metadata.get("align"), "center")

    def test_run_underline_from_character_style_is_preserved(self):
        """run 未直接设置 underline，靠引用的字符样式（w:rStyle）带下划线。"""
        from docx import Document as DocxDocument
        from docx.enum.style import WD_STYLE_TYPE

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc = DocxDocument()
            char_style = doc.styles.add_style("QuietLink", WD_STYLE_TYPE.CHARACTER)
            char_style.font.underline = True
            para = doc.add_paragraph()
            run = para.add_run("styled underline text")
            run.style = char_style
            path = os.path.join(tmp_dir, "styled-underline.docx")
            doc.save(path)

            result = DocxParser()._parse_docx(path)
            chunks = [c for page in result.pages for c in page.chunks]
            paragraph_chunk = next(c for c in chunks if c.chunk_type == "paragraph")
            runs_meta = paragraph_chunk.metadata.get("runs") or []
            self.assertTrue(runs_meta)
            self.assertTrue(runs_meta[0].get("underline"))

    def test_run_underline_from_paragraph_style_default_run_props_is_preserved(self):
        """run 和其字符样式都没提下划线，靠段落样式的默认 run 属性继承。"""
        from docx import Document as DocxDocument

        with tempfile.TemporaryDirectory() as tmp_dir:
            doc = DocxDocument()
            normal_style = doc.styles["Normal"]
            normal_style.font.underline = True
            para = doc.add_paragraph()
            para.add_run("plain text inheriting paragraph style underline")
            path = os.path.join(tmp_dir, "styled-paragraph-underline.docx")
            doc.save(path)

            result = DocxParser()._parse_docx(path)
            chunks = [c for page in result.pages for c in page.chunks]
            paragraph_chunk = next(c for c in chunks if c.chunk_type == "paragraph")
            runs_meta = paragraph_chunk.metadata.get("runs") or []
            self.assertTrue(runs_meta)
            self.assertTrue(runs_meta[0].get("underline"))


if __name__ == "__main__":
    unittest.main()
